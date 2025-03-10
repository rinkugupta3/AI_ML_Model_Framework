# NLP Model uses spaCy library as small "en_core_web_sm"
# Install pip install python-docx
# Install pip install spacy
# Install python -m spacy download en_core_web_sm
# pytest -s -v tests/test_ai_nlp_model.py

import os
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import docx  # Import docx
import pytest
import re  # Import the re module

def generate_test_plan_docx(user_story_file=None,
                             user_story_text=None,
                             output_file="test_plan.docx"):
    """
    Generates a basic test plan document in .docx format.  It can read the user
    story from a .docx file, or you can pass in the user story text directly.
    """

    as_a = ""
    i_want_to = ""
    so_that = ""
    acceptance_criteria_list = []
    scenarios_covered_list = []

    try:
        if user_story_file:
            # Read the user story content from the .docx file
            doc = docx.Document(user_story_file)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            user_story_text = '\n'.join(full_text)

        if not user_story_text:
            raise ValueError("No user story provided (either file or text).")

        # --- Extract the components (use more flexible regex) ---
        try:
            user_story_match = re.search(r"As a (.*), I want to (.*) So that (.*)", user_story_text, re.DOTALL)
            if user_story_match:
                as_a = user_story_match.group(1).strip()
                i_want_to = user_story_match.group(2).strip()
                so_that = user_story_match.group(3).strip()
            else:
                print("Warning: Could not parse user story using regex.  Check the format.")
                as_a = "Could not parse: Check format"
                i_want_to = "Could not parse: Check format"
                so_that = "Could not parse: Check format"
        except Exception as e:
            print(f"Error extracting user story components: {e}")
            as_a = "Error parsing"
            i_want_to = "Error parsing"
            so_that = "Error parsing"


        # Extract Acceptance Criteria and Scenarios (more flexible extraction)
        try:
            acceptance_criteria_match = re.search(r"(?i)(Acceptance Criteria:|Acceptance Criteria:)\s*([\s\S]*?)(?=(Scenarios Covered:|Test Scenarios:|\Z))",
                                                  user_story_text, re.DOTALL)
            if acceptance_criteria_match:
                acceptance_criteria_text = acceptance_criteria_match.group(2).strip()
                acceptance_criteria_list = [s.strip() for s in re.split(r"\n*\d+\.\s*", acceptance_criteria_text) if s.strip()] #splits the text on newline followed by a number and a period
            else:
                acceptance_criteria_list = ["No Acceptance Criteria found or could not parse"]
        except Exception as e:
            print(f"Error extracting acceptance criteria: {e}")
            acceptance_criteria_list = ["Error parsing acceptance criteria"]

        try:
            scenarios_covered_match = re.search(r"(?i)(Scenarios Covered:|Test Scenarios:)\s*([\s\S]*)",
                                                 user_story_text, re.DOTALL)
            if scenarios_covered_match:
                scenarios_covered_text = scenarios_covered_match.group(2).strip()
                scenarios_covered_list = [s.strip() for s in re.split(r"\n*\d+\.\s*", scenarios_covered_text) if s.strip()]
            else:
                scenarios_covered_list = ["No Scenarios Covered found or could not parse"]
        except Exception as e:
            print(f"Error extracting scenarios covered: {e}")
            scenarios_covered_list = ["Error parsing scenarios"]


    except Exception as e:
        print(f"General error processing user story: {e}")
        as_a = "General Parsing Error"
        i_want_to = "General Parsing Error"
        so_that = "General Parsing Error"
        acceptance_criteria_list = ["General Parsing Error"]
        scenarios_covered_list = ["General Parsing Error"]

    # Create a new Document
    document = Document()

    # Styles
    styles = document.styles

    # Heading Style
    heading1_style = styles.add_style('Heading1Style', WD_STYLE_TYPE.PARAGRAPH)
    heading1_font = heading1_style.font
    heading1_font.name = 'Arial'
    heading1_font.size = Pt(16)

    heading2_style = styles.add_style('Heading2Style', WD_STYLE_TYPE.PARAGRAPH)
    heading2_font = heading2_style.font
    heading2_font.name = 'Arial'
    heading2_font.size = Pt(14)

    # List Paragraph Style
    list_paragraph_style = styles.add_style('ListParagraph', WD_STYLE_TYPE.PARAGRAPH)
    list_paragraph_font = list_paragraph_style.font
    list_paragraph_font.name = 'Arial'
    list_paragraph_font.size = Pt(12)

    paragraph_style = styles.add_style('NormalStyle', WD_STYLE_TYPE.PARAGRAPH)
    paragraph_font = paragraph_style.font
    paragraph_font.name = 'Arial'
    paragraph_font.size = Pt(12)


    # Add a heading
    document.add_heading('Test Plan: Login and Logout Functionality', level=1)
    document.paragraphs[-1].style = document.styles['Heading1Style']

    # Add the generation timestamp
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    p = document.add_paragraph(f'Generated on: {timestamp}')
    p.style = document.styles['NormalStyle']  # Apply the style


    # Add Introduction section
    document.add_heading('1. Introduction', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']

    document.add_paragraph('Purpose: This test plan outlines the strategy for testing the login and logout functionality of the application to ensure consistent user access control, security, and cross-browser compatibility.')
    document.paragraphs[-1].style = document.styles['NormalStyle']

    document.add_paragraph('Scope: This test plan covers the login and logout functionality across Chromium, Firefox, and WebKit browsers, using valid and invalid user credentials as defined in `Config.USERS`.')
    document.paragraphs[-1].style = document.styles['NormalStyle']

    # Add User Story section
    document.add_heading('2. User Story', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']

    document.add_paragraph(f"**As a:** {as_a}")
    document.paragraphs[-1].style = document.styles['NormalStyle']

    document.add_paragraph(f"**I want to:** {i_want_to}")
    document.paragraphs[-1].style = document.styles['NormalStyle']

    document.add_paragraph(f"**So that:** {so_that}")
    document.paragraphs[-1].style = document.styles['NormalStyle']

    # Add Acceptance Criteria section
    document.add_heading('3. Acceptance Criteria', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']

    for criteria in acceptance_criteria_list:
        p = document.add_paragraph(f"{criteria.strip()}", style='ListParagraph')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25) #Hangind indent

    # Add Scenarios Covered section
    document.add_heading('4. Scenarios Covered', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']

    for scenario in scenarios_covered_list:
        p = document.add_paragraph(f"{scenario.strip()}", style='ListParagraph')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    # Add Test Cases section with a table
    document.add_heading('5. Test Cases', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']
    table = document.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1'  # Optional table style
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test Case ID'
    hdr_cells[1].text = 'Test Description'
    hdr_cells[2].text = 'Preconditions'
    hdr_cells[3].text = 'Test Steps'
    hdr_cells[4].text = 'Expected Results'
    hdr_cells[5].text = 'Priority'

    # Add test cases
    test_cases = [
        ("TC_LOGIN_001", "Successful Login & Logout Flow (Valid Credentials)", "`Config.USERS` contains valid username and password. App is running.", "1. Navigate to `Config.BASE_URL`. 2. Enter valid credentials. 3. Click login. 4. Verify dashboard. 5. Click logout. 6. Verify login page.", "User logs in, sees dashboard, logs out, sees login page.", "High"),
        ("TC_LOGIN_002", "Failed Login Attempt (Invalid Credentials)", "`Config.USERS` contains invalid credentials. App is running.", "1. Navigate to `Config.BASE_URL`. 2. Enter invalid credentials. 3. Click login.", "User remains on login page. Error message displayed (if any).", "High"),
        ("TC_LOGIN_003", "Login with Empty Credentials", "App is running.", "1. Navigate to `Config.BASE_URL`. 2. Attempt login with empty fields. 3. Click login.", "User remains on login page. Error message displayed (if any).", "Medium")
    ]

    for test_case_id, test_description, preconditions, test_steps, expected_results, priority in test_cases:
        row_cells = table.add_row().cells
        row_cells[0].text = test_case_id
        row_cells[1].text = test_description
        row_cells[2].text = preconditions
        row_cells[3].text = test_steps
        row_cells[4].text = expected_results
        row_cells[5].text = priority

    # Add Test Environment section
    document.add_heading('6. Test Environment', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']

    document.add_paragraph('Operating System: Tests will be executed on a platform that supports Chromium, Firefox, and WebKit browsers (e.g., Windows, macOS, Linux).')
    document.paragraphs[-1].style = document.styles['NormalStyle']

    p = document.add_paragraph()
    p.style = document.styles['ListParagraph']
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.add_run("Browsers:").bold = True

    document.add_paragraph("Chromium (latest stable version)",style='ListParagraph')
    p = document.add_paragraph("Firefox (latest stable version)",style='ListParagraph')

    document.add_paragraph("WebKit (latest version)", style='ListParagraph')


    document.add_paragraph('Test Framework: Playwright with Python.')
    document.paragraphs[-1].style = document.styles['NormalStyle']

    document.add_paragraph('Configuration: `Config.USERS` must be accessible to the test framework. Ensure `Config.BASE_URL` is correctly set.')
    document.paragraphs[-1].style = document.styles['NormalStyle']

    document.add_paragraph('CI/CD: Tests should be executable within a CI/CD pipeline.')
    document.paragraphs[-1].style = document.styles['NormalStyle']

   # Add Test Data section
    document.add_heading('7. Test Data', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']

    document.add_paragraph('Users: The test data will be sourced from `Config.USERS`.')
    document.paragraphs[-1].style = document.styles['NormalStyle']


    p = document.add_paragraph()
    p.style = document.styles['ListParagraph']
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.add_run("Valid Users:").bold = True

    document.add_paragraph('At least one user with valid credentials.', style='ListParagraph')


    p = document.add_paragraph()
    p.style = document.styles['ListParagraph']
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.add_run("Invalid Users (Optional):").bold = True

    document.add_paragraph('Users with invalid credentials.', style='ListParagraph')

    document.add_paragraph('`Config.BASE_URL`: Ensure this is set correctly to the application\'s login page URL.')
    document.paragraphs[-1].style = document.styles['NormalStyle']

    # Add Test Execution section
    document.add_heading('8. Test Execution', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']

    document.add_paragraph('Tests will be executed using the Playwright test framework.')
    document.paragraphs[-1].style = document.styles['NormalStyle']

    # Add Test Deliverables section
    document.add_heading('9. Test Deliverables', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']

    deliverables = ["Automated test scripts", "Test reports", "Screenshots", "Test execution logs"]
    for deliverable in deliverables:
        p = document.add_paragraph(deliverable, style='ListParagraph')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)



    # Add Entry and Exit Criteria section
    document.add_heading('10. Entry and Exit Criteria', level=2)
    document.paragraphs[-1].style = document.styles['Heading2Style']

    document.add_paragraph('Entry Criteria:', style='NormalStyle')
    entry_criteria = ["Application is deployed.", "Test environment is set up.", "Test data is available.", "Test scripts are validated."]
    for item in entry_criteria:
         p = document.add_paragraph(item, style='ListParagraph')
         p.paragraph_format.left_indent = Inches(0.25)
         p.paragraph_format.first_line_indent = Inches(-0.25)


    document.add_paragraph('Exit Criteria:', style='NormalStyle')
    exit_criteria = ["All test cases executed.", "Test reports generated and analyzed.", "All critical defects resolved.", "Login/logout meets criteria."]
    for item in exit_criteria:
        p = document.add_paragraph(item, style='ListParagraph')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)


    # Ensure the "documents" directory exists
    documents_dir = "documents"
    if not os.path.exists(documents_dir):
        os.makedirs(documents_dir)

    output_path = os.path.join(documents_dir, output_file)

    try:
        document.save(output_path)  # Save the document
        print(f"Test plan generated successfully at: {output_path}")
    except Exception as e:
        print(f"Error generating test plan: {e}")

# Example Usage: Provide either the user story file OR the user story text
user_story_file = r"C:\Users\dhira\Desktop\Dhiraj HP Laptop\Projects\AI_Model_Driven_TestCases_Automation_Script\Login and Logout Functionality Validation Across Multiple Browsers.docx"  # **COMPLETE PATH HERE**
# If you don't want to read from a file, comment out the above line and uncomment the lines below
#user_story_text = """
#As a Quality Assurance Engineer,
#I want to verify the login and logout functionality of the application across multiple browsers and user accounts,
#So that I can ensure consistent user access control, security, and cross-browser compatibility.
#
#Acceptance Criteria:
#1. Cross-Browser Testing: Validate login/logout flow on Chromium, Firefox, and WebKit browsers.
#2. Multi-User Credential Validation: Test all configured users (valid and invalid credentials) in Config.USERS.
#3. Post-Login Redirection: After successful login, users must be redirected to the dashboard.
#4. Logout Functionality: Logout must redirect users to the login page (Config.BASE_URL).
#
#Scenarios Covered:
#1. Successful Login & Logout Flow: Navigate to the login page. Enter valid credentials. Confirm redirection to the dashboard. Log out and verify redirection to the login page.
#2. Failed Login Handling: Attempt login with invalid credentials (if applicable). Verify the user remains on the login page.
#"""

generate_test_plan_docx(user_story_file=user_story_file)  # Pass the file